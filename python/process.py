import pprint
import docx
import re
import os
import openpyxl
import datetime
import time


''' Конфигурация '''
xlsx_file = './xlsx/PersonalInfo.xlsx'
templates_dir = './templates'
out_dir = './out'
time_profile = None


''' Функции для профайлинга '''
def start_time():
	global time_profile
	time_profile = time.time()


def stop_time():
	global time_profile
	diff = time.time() - time_profile
	time_profile = None
	print('[%.3f сек.]' % diff)


''' Список паттернов для сопоставления номеров групп и специальности '''
patterns = [
	{ 'pattern': r'3.$', 				'spec': r'Акушерское дело' },
	{ 'pattern': r'4.$', 				'spec': r'Лабораторная диагностика' },
	{ 'pattern': r'5.$', 				'spec': r'Стоматология ортопедическая' },
	{ 'pattern': r'6.$', 				'spec': r'Фармация' },
	{ 'pattern': r'8.$', 				'spec': r'Медико-профилактическое дело' },
	{ 'pattern': r'9.$', 				'spec': r'Социальная работа' },
	{ 'pattern': r'^(2|0).(1|2).$', 	'spec': r'Сестринское дело' },
	{ 'pattern': r'^1.1.$', 			'spec': r'Младшая медицинская сестра' },
	{ 'pattern': r'^1.2.$', 			'spec': r'Лечебное дело' }
]


def read_xlsx(filename: 'str') -> '[{group, f, i}]':
	''' Функция для чтения файлов с группами и ФИО из АСУ РСО '''
	workbook = openpyxl.load_workbook(filename = filename)
	worksheet = workbook.active
	students = []
	group = ''
	for row in worksheet.rows:
		''' Определяем группу '''
		if type(row[0].value) == str:
			if re.search(r'Группа', row[0].value):
				group = re.sub(r'Группа ', '', row[0].value)
		''' Выделяем ФИО (если напротив - дата рождения) '''
		if type(row[3].value) == datetime.datetime:
			fio = re.sub(r'^(\S+)\s+(\S+).*', r'\1 \2', row[1].value)
			fio = re.split(r'\s+', fio)
			f = re.sub(r'^\s+|\s+$', '', fio[0])
			i = re.sub(r'^\s+|\s+$', '', fio[1])
			students.append({ 'group': group, 'f': f, 'i': i })
	return students


def get_spec_hash(students: '[{group, f, i}]') -> 'dict(spec => [{group, f, i}])':
	''' Функция сортирует студентов по специальностям '''
	global patterns # Паттерны для определения групп
	hsh = dict()
	for student in students:
		# Находим специальность
		group = student['group']
		spec = None
		for pt in patterns:
			if re.search(pt['pattern'], group):
				spec = pt['spec']
		# Если не находим, то ошибка в паттернах
		if spec == None:
			raise Exception('Ошибка! Группа %s не отнесена ни к одной специальности!' % group)
		# Создаем хэш, где ключ - специальность
		if not(spec in hsh):
			hsh[spec] = []
		# Добавляем студента в хэш
		(hsh[spec]).append(student)
	return hsh


def fill_templates(students_hsh: 'dict(spec => [{group, f, i}])') -> 'None':
	''' Функция заполняет каждый шаблон студентами '''
	global out_dir
	global templates_dir
	tpl_files = os.listdir(templates_dir)
	# Обрабатываем списки студентов по каждой специальности
	for spec in students_hsh:
		print(' --> Заполняем специальность %s (%s чел.)' % (spec, len(students_hsh[spec])))
		# Открываем файл шаблона или выдаем ошибку
		template = list(filter(lambda x: re.search(spec, x), tpl_files))
		if (len(template) == 0):
			raise Exception('Ошибка! Файл шаблона для специальности %s не найден!' % spec)
		# Пути к файлу шаблона и выходному файлу 
		in_template = '%s/%s' % (templates_dir, template[0])
		out_template = '%s/%s' % (out_dir, template[0])

		doc = docx.Document(in_template)
		# Заполняем таблицу в шаблоне
		table = doc.tables[0]
		for row in table.rows:
			for cell in row.cells:
				# Если есть в пуле студент - выбираем его
				student = None
				if len(students_hsh[spec]) > 0:
					student = students_hsh[spec].pop(0)
				# Заполняем ячейку данными студента
				if student == None:
					cell.text = ''
				else:
					for paragraph in cell.paragraphs:
						for run in paragraph.runs:
							if re.search(r'Фамилия', run.text):
								text = run.text.replace(r'Фамилия', student['f'])
								run.text = text
							if re.search(r'Имя', run.text):
								text = run.text.replace(r'Имя', student['i'])
								run.text = text
							if re.search(r'Нгруппа', run.text):
								text = run.text.replace(r'Нгруппа', student['group'])
								run.text = text
		# Если не все студенты внесены, выводим ошибку
		if len(students_hsh[spec]) > 0:
			raise Exception('Ошибка! В шаблоне не хватает ячеек для %s человек' % len(students_hsh[spec]))
		# Сохраняем изменения
		doc.save(out_template)




if __name__ == '__main__':
	print('1. Читаем файл из АСУ РСО')
	start_time()
	students = read_xlsx(xlsx_file)
	print(' --> Количество студентов: %s чел.' % len(students))
	stop_time()

	print('2. Распределяем по специальностям')
	start_time()
	students_hsh = get_spec_hash(students)
	stop_time()

	print('3. Заполняем шаблоны')
	start_time()
	fill_templates(students_hsh)
	stop_time()
